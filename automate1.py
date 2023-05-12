from docx import Document
from docx.opc.exceptions import PackageNotFoundError


def get_para_data(output_doc_name, paragraph):
    try:
        output_para = output_doc_name.add_paragraph()
        for run in paragraph.runs:
            output_run = output_para.add_run(run.text)
            output_run.bold = run.bold

            output_run.italic = run.italic
            output_run.underline = run.underline
            output_run.font.color.rgb = run.font.color.rgb

            output_run.style.name = run.style.name
        output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
    except PackageNotFoundError:
        print('error')

input_doc = Document(input('enter the doc file: '))
output_doc = Document()

for para in input_doc.paragraphs:
    get_para_data(output_doc, para)

output_doc.save('hello.ppt')
print('saved as hello.ppt')